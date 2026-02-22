Entiendo perfectamente el problema. Las capturas muestran que el sistema está bloqueando la visualización del código dentro de los cuadros grises, dejándote solo con el texto explicativo. Como tu perfil no es técnico, esto hace que sea imposible copiar lo que necesitas.

Para resolver esto, voy a enviarte el código sin usar cuadros de código. Lo escribiré como texto plano directamente aquí abajo. Solo tienes que seleccionar todo el texto desde donde empieza import streamlit hasta el final, y pegarlo en tu archivo app.py.

CÓDIGO PARA APP.PY (COPIAR DESDE AQUÍ ABAJO)
import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
import plotly.graph_objects as go
import json
from fpdf import FPDF
from docx import Document
from io import BytesIO

Configuración del modelo y API
try:
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
model = genai.GenerativeModel(
model_name='gemini-1.5-flash',
generation_config={"response_mime_type": "application/json"}
)
except Exception as e:
st.error(f"Error de configuración: {e}")

st.set_page_config(page_title="Auditor de Talento IA", layout="wide")

Funciones de exportación de reportes
def generar_pdf(resultados):
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", "B", 16)
pdf.cell(190, 10, "Reporte de Analisis de CVs", ln=True, align='C')
for res in resultados:
pdf.ln(10)
pdf.set_font("Arial", "B", 12)
nombre = str(res.get('nombre_candidato', 'N/A')).encode('latin-1', 'ignore').decode('latin-1')
pdf.cell(190, 10, f"Candidato: {nombre}", ln=True)
pdf.set_font("Arial", "", 10)
resumen = str(res.get('resumen_reservas', '')).encode('latin-1', 'ignore').decode('latin-1')
pdf.multi_cell(190, 7, f"Resumen: {resumen}")
pdf.cell(190, 7, f"Score: {res.get('experiencia')}% Exp | {res.get('habilidades')}% Hab", ln=True)
return pdf.output(dest='S').encode('latin-1', errors='replace')

def generar_docx(resultados):
doc = Document()
doc.add_heading('Evaluacion de Candidatos', 0)
for res in resultados:
doc.add_heading(f"Candidato: {res.get('nombre_candidato', 'N/A')}", level=1)
doc.add_paragraph(f"Nota IA: {res.get('resumen_reservas', '')}")
doc.add_paragraph(f"Puntajes: Exp: {res.get('experiencia')}% | Edu: {res.get('educacion')}% | Hab: {res.get('habilidades')}%")
bio = BytesIO()
doc.save(bio)
return bio.getvalue()

Interfaz de la aplicación
st.title("🛡️ Plataforma de Evaluación de Candidatos")

if "autorizado" not in st.session_state:
st.session_state.autorizado = False

if not st.session_state.autorizado:
st.warning("### Aviso de Privacidad")
st.write("Al continuar, usted autoriza el procesamiento de datos personales contenidos en los CVs.")
if st.button("Acepto los términos"):
st.session_state.autorizado = True
st.rerun()
else:
st.sidebar.header("Vacante")
jd_text = st.sidebar.text_area("Descripción del Puesto:", height=200)
uploaded_files = st.file_uploader("Cargar CVs (PDF)", type="pdf", accept_multiple_files=True)
